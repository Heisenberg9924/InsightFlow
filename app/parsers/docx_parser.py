from pathlib import Path
from uuid import uuid4

from docx import Document

from app.models.document import ParsedDocument, DocumentMetadata


def extract_docx_text(path: str) -> ParsedDocument:

    doc = Document(path)

    text = "\n".join(p.text for p in doc.paragraphs)

    metadata = DocumentMetadata(
        word_count=len(text.split())
    )

    return ParsedDocument(
        id=str(uuid4()),
        filename=Path(path).name,
        file_type="docx",
        text=text,
        metadata=metadata
    )